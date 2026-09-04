#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Build the WSQ assessment set for 'Applications Integration with Power Apps and Power Automate'
(TGS-2022015539):
  - Written Assessment (SAQ)  — 5 open-ended KNOWLEDGE questions (K1–K5), aligned to the slides
  - Practical Performance (PP) — 4 PRACTICAL tasks (A1–A8), aligned to the in-class labs
Mirrors the original paper held on the TMS: same instrument, same question/task count, same K/A
codes and mapping, same timings (WA 1 hour, PP 90 minutes). Only the content is rewritten, from
this course's slides and labs.
Each instrument is produced as a Question Paper and a matching Answer Key (4 DOCX total),
all with the WSQ house cover page (same as the Lesson Plan / Learner Guide). Page 1 is the cover;
page 2 carries Trainee Information + Instructions + Grading; the questions/tasks begin on page 3.
Body: Arial 11.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# This script lives in the wsq-assessment skill (.claude/skills/wsq-assessment/) and runs in
# place — it detects the course repo root by walking up to the nearest dir that has a .git
# folder (or both courseware/ and assessment/). Override with env REPO=/path if needed.
def _find_repo():
    env = os.environ.get("REPO")
    if env and os.path.isdir(env):
        return os.path.abspath(env)
    d = os.path.dirname(os.path.abspath(__file__))
    while d != os.path.dirname(d):
        if os.path.isdir(os.path.join(d, ".git")) or \
           (os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "assessment"))):
            return d
        d = os.path.dirname(d)
    return os.getcwd()

REPO = _find_repo()
# prodoc.py (WSQ cover page + version control + page numbers, same as LP/LG) ships with the
# tertiary-lesson-plan skill. Look for it at the project level first, then the user level.
for _cand in (os.path.join(REPO, ".claude/skills/tertiary-lesson-plan"),
              os.path.expanduser("~/.claude/skills/tertiary-lesson-plan")):
    if os.path.exists(os.path.join(_cand, "prodoc.py")):
        sys.path.insert(0, _cand); break
import prodoc  # cover page + version control + page numbers (same as LP/LG)

# ─── EDIT PER COURSE ────────────────────────────────────────────────────────
TITLE       = "Applications Integration with Power Apps and Power Automate"
COURSE_CODE = "TGS-2022015539"
# ────────────────────────────────────────────────────────────────────────────
# The cover page renders prodoc's module-level TGS constant. Override it so the
# assessment cover shows THIS course's ref (works with either prodoc version —
# the older project prodoc has no course_code kwarg).
prodoc.TGS = f"TGS Ref No: {COURSE_CODE}"
OUT   = os.path.join(REPO, "assessment")

# Logos: prefer the course's own courseware/assets, else fall back to the copies bundled
# in this skill (so the assessment builds even outside this project). Replace the course
# logo per course; the Tertiary Infotech logo is the same for every WSQ course.
def _logo(name):
    here = os.path.dirname(os.path.abspath(__file__))
    for p in (os.path.join(REPO, "courseware/assets", name), os.path.join(here, "assets", name)):
        if os.path.exists(p):
            return p
    return None
ORG_LOGO    = _logo("tertiary-infotech-logo.png")
COURSE_LOGO = _logo("power-platform-course-logo.png")   # None if absent → Tertiary-only cover (as LP/LG)

Q_VER, A_VER = "v4", "v4"   # single standardised version across all four files
BRAND = RGBColor(0x1F, 0x6F, 0xEB); DARK = RGBColor(0x11, 0x18, 0x27); GREY = RGBColor(0x55, 0x5B, 0x66)
# Assessments carry the cover page only — no Document Version Control Record.

# ---------------------------------------------------------------- WRITTEN (KNOWLEDGE)
# (criterion, context, question, [model-answer points]) — each traces to the course slides.
# Mirrors the original paper: 5 open-ended questions, K1-K5, 1 hour, open book.
WRITTEN = [
 ("K1",
  "Your manager wants to reduce the manual re-keying that happens between the service desk, the "
  "engineering team and the monthly report. Before anything is built, the opportunities have to be "
  "identified and the right kind of middleware chosen.",
  "Your manager wants to automate routine tasks across multiple departments, such as logging service "
  "calls, setting reminders and managing approval processes. Using Power Automate, list three types of "
  "applications or tasks that could be automated to streamline operations, and explain what makes "
  "Power Automate a suitable type of middleware for them.",
  ["Three candidate automations, for example: (1) service-call intake — an incoming email or form "
   "response automatically logged as a row in Excel or a SharePoint list; (2) approval processes — "
   "leave, purchase or expense requests routed to an approver and the decision written back; "
   "(3) scheduled reporting or reminders — a daily or weekly digest assembled from a data source and "
   "e-mailed to a distribution list.",
   "Other acceptable answers: file synchronisation between OneDrive/SharePoint and another store; "
   "notification of new records to a Teams channel; data collection from Microsoft Forms into a table.",
   "Power Automate is an integration platform (iPaaS) — cloud middleware that brokers the exchange "
   "between applications so no human retypes the data.",
   "Its features make it suitable: 1,000+ prebuilt connectors covering Microsoft 365 and third-party "
   "SaaS; custom connectors for any REST API; low-code so business users can build and maintain flows; "
   "triggers for event, manual and scheduled starts; and desktop flows (RPA) to reach legacy systems "
   "that expose no API.",
   "Each candidate should be justified by the volume of transactions, the error rate of doing it by "
   "hand, and the business impact of the delay it removes."]),

 ("K2",
  "A colleague is building his first flow and is stuck at the very first decision — what should start "
  "the flow, and how do actions receive the data the trigger produced.",
  "Tom is trying to create a flow in Power Automate to automatically copy files between two cloud "
  "services, but he is unsure how to find the right trigger to start the flow. Explain how Tom should "
  "proceed to find the most suitable trigger, and how the file data then reaches the later actions.",
  ["A trigger is the single event that starts a cloud flow; every flow has exactly one, and it is "
   "always the first block in the flow.",
   "Tom should first decide the flow TYPE: this is event-driven, so an Automated cloud flow is correct "
   "(instant = started by a person, scheduled = started on a recurrence, desktop = RPA).",
   "He should then search for the SOURCE service by name in the trigger search box — for example "
   "OneDrive for Business — and choose the trigger that matches the event he wants, 'When a file is "
   "created'.",
   "He can browse the connector catalogue at make.powerautomate.com/connectors to confirm the "
   "connector exists and whether it is Standard or Premium, since Premium changes the licence needed.",
   "He should check the trigger's own settings and Advanced options — folder, filter, include-contents "
   "— to narrow what actually fires the flow rather than filtering later, which is cheaper and clearer.",
   "Data reaches the later actions through DYNAMIC CONTENT: the trigger's outputs (file name, file "
   "content, identifier) are offered as tokens which Tom inserts into the inputs of the destination "
   "'Create file' action. This is how data crosses the boundary between the two applications.",
   "He should then use the Flow checker to clear errors, Test the flow, and confirm the run succeeded "
   "in the 28-day run history."]),

 ("K3",
  "Applications run on many different platforms, and a Power App can bind to data held in very "
  "different places. A colleague's app cannot see the Excel data she wants to use.",
  "Your colleague is using an Excel workbook as a data source for their Power App, but the app is not "
  "recognizing the data. What should your colleague do to make the data usable in Power Apps, and what "
  "does this tell you about how apps connect to data across platforms?",
  ["The workbook data must be formatted as a real Excel TABLE. Power Apps and Power Automate can only "
   "bind to a named table, never to a loose range of cells — this is the most common cause of the "
   "symptom described.",
   "Fix: open the workbook, select the data including the header row, choose Insert > Table with 'My "
   "table has headers' ticked, then give the table a clear name in Table Design (for example "
   "ServiceCalls).",
   "The workbook must be stored in a cloud location the connector can reach — OneDrive for Business, "
   "SharePoint, Google Drive or Dropbox — not on the local machine.",
   "The file must be .xlsx and should be closed, since an open or locked file can block the connector.",
   "In the app, add the data source through the correct connector (Excel Online (Business)) and select "
   "the workbook and the table; then refresh the data source.",
   "What this shows about platforms: an app never talks to a file directly — it reaches every data "
   "source through a CONNECTOR, which is a packaged API. The same canvas app can therefore bind to "
   "Dataverse, SharePoint, SQL Server, Excel or a custom REST API, and the app itself runs across "
   "browser, iOS, Android and embedded in Teams.",
   "Excel is convenient for teaching and small volumes, but has real limits — weak concurrency, no "
   "row-level security and poor delegation — so SharePoint or Dataverse is the better production "
   "choice."]),

 ("K4",
  "After an integration is built, issues appear that were not visible at design time. Recognising the "
  "class of issue is what tells you where to look for the fix.",
  "You are customizing a three-screen canvas app and integrating it with a Power Automate flow. "
  "Describe the potential technical, compatibility and performance issues you should be aware of, and "
  "give one concrete example and mitigation for each class.",
  ["Integration issues fall into three classes — technical, compatibility and performance — and naming "
   "the class tells you where to look.",
   "TECHNICAL — the connection itself. Example: the flow fails with a 401/403 because a connection has "
   "expired, was revoked, or was never authenticated; or a DLP policy separates two connectors the flow "
   "combines; or a Premium connector is used without the licence. Mitigation: re-authenticate the "
   "connection and save; run business-critical flows under a dedicated service account rather than a "
   "personal one; check the environment's DLP policies before designing across connector groups.",
   "COMPATIBILITY — the shape of the data. Example: the app passes Days as text where the flow's "
   "Power Apps (V2) trigger declares a number, so the call fails schema validation; or a date arrives "
   "in a format the target cannot parse; or a renamed Excel column breaks the mapping. Mitigation: "
   "coerce the type at the call site, for example Value(txtDays.Text) instead of txtDays.Text; treat "
   "connector inputs as a typed contract; keep column names stable and re-point actions after a rename.",
   "PERFORMANCE — scale and throughput. Example: a gallery uses a non-delegable function so only the "
   "first 500 rows are processed and the app silently shows an incomplete result; or the connector "
   "returns a throttling error under load. Mitigation: rewrite to a delegable function such as "
   "StartsWith() or Filter() on an indexed column until the blue delegation warning clears; reduce rows "
   "at source; raise the data row limit in Settings > General as a stop-gap only; batch high-volume calls.",
   "On the app's appearance specifically: a canvas app has no global theme applied automatically, so a "
   "colour or font changed on one screen does NOT propagate to the others. Set the styling on each "
   "screen, or drive it from global variables or a component library so one change updates every "
   "screen consistently.",
   "Delegation is the issue to watch most closely, because the app does not fail — it returns a wrong "
   "answer quietly, which is more dangerous than an error."]),

 ("K5",
  "An API is how one application exposes its functions to another. In the Power Platform, connectors "
  "are the API surface, and apps can also be surfaced inside other platforms such as Microsoft Teams.",
  "You have built a Help Desk app using Power Apps, and the staff primarily work within Microsoft "
  "Teams. Explain the function of an API (and of connectors) in this integration, and describe the "
  "steps you would take to make the app available to staff inside Teams.",
  ["An API (Application Programming Interface) is the defined contract through which one application "
   "exposes its data and functions to another — the request format, the response schema, and the "
   "authentication required. It lets systems interoperate without knowing each other's internals.",
   "In the Power Platform, a CONNECTOR is a packaged API: it wraps a service's REST API and presents "
   "its operations as triggers and actions with typed inputs and outputs. Standard connectors are "
   "included in most licences, Premium connectors need a higher plan, and a custom connector can be "
   "built over any reachable REST API by defining its host, base URL, security and actions.",
   "Steps to surface the app in Teams: first save and PUBLISH the app in Power Apps, since only a "
   "published version is served to users.",
   "Share the app with the staff or the security group who need it, granting the User role (not "
   "Co-owner) so they can run it but not edit it.",
   "Confirm those users also have access to the underlying DATA SOURCE — sharing an app does not share "
   "its data, and this is a common reason a shared app fails for everyone but the maker.",
   "In Microsoft Teams, add the Power Apps app to the workspace, then choose 'Add an app to a channel "
   "tab', select the Help Desk app and add it as a tab so it appears in the channel the staff use.",
   "Alternatively use Apps > Built by your colleagues in Teams to find the shared app, or embed it via "
   "the Power Apps tab in a chat or meeting.",
   "Verify by opening the tab as one of the staff users and completing an end-to-end action, checking "
   "that the app renders correctly in the Teams client on both desktop and mobile."]),
]

# ---------------------------------------------------------------- PRACTICAL (ACTIVITY-BASED)
SCENARIO = (
 "You have been assigned to create an Employee Engagement Survey solution using Power Apps and Power "
 "Automate as part of your organisation's initiative to gather feedback from employees. The app needs "
 "to be connected to a cloud data source and automated to notify the relevant stakeholders whenever a "
 "survey is submitted. After confirming that the solution works as expected, you will save, publish "
 "and share it with your assessor for review. Throughout the process you will identify and address any "
 "technical, compatibility or performance issues so that the solution is stable and usable. "
 "Complete the four tasks below; each mirrors a hands-on lab you did in class. Build everything inside "
 "the course Power Platform environment, and paste your formulas, flow configuration and screenshots "
 "as evidence.")

# (label, criterion, task prompt, box caption, model-answer build steps citing the lab)
BOX_CAP = "Paste your formulas / flow configuration and screenshots of your output in the box below"
PRACTICAL = [
 ("Task 1", "A1, A2",
  "Identify the integration opportunity and assess its feasibility, then create the app. "
  "Part A — For the Employee Engagement Survey scenario, identify at least THREE opportunities for "
  "creating connections among the applications involved (for example the survey front end, the data "
  "store, the notification channel and the reporting step). For each, name the source system, the "
  "target system and the connector you would use. "
  "Part B — Perform a feasibility scan on your chosen design against the six criteria taught in class "
  "(connector availability, data readiness, licensing, security and governance, performance, "
  "maintainability), scoring each and recording a Build / Defer / Reject decision with a justification. "
  "Part C — Create a new canvas app in the course environment for the survey and customise its "
  "appearance — the app name, icon, colour scheme and text elements — to match your organisation's "
  "branding. (Labs 1-2 — Integration Opportunity Scan and Feasibility Scan; Lab 7-8 — building a "
  "canvas app.)",
  BOX_CAP,
  "Part A — Integration Opportunity Register, at least three rows, each naming source, target and connector:\n"
  "  1. Survey front end (canvas app) -> data store (Excel Online (Business) / SharePoint / Dataverse connector)\n"
  "  2. Data store -> notification (Office 365 Outlook connector sends an e-mail on submission)\n"
  "  3. Data store -> approval or review (Approvals / Microsoft Teams connector routes to a stakeholder)\n"
  "  4. (optional) Data store -> reporting (Power BI connector, or a scheduled digest flow)\n"
  "  Each row classified by what is being joined — device, database, software or application — and scored\n"
  "  1-5 for volume, error rate and business impact, then ranked.\n"
  "Part B — Feasibility Assessment Matrix scoring the chosen design 1-5 on each of the six criteria:\n"
  "  Connector availability — a standard connector exists for Excel/SharePoint and Outlook; note Standard vs Premium\n"
  "  Data readiness — the workbook data is a real Excel TABLE with a stable header row and a clear key\n"
  "  Licensing — Standard connectors are covered by the Microsoft 365 plan; Premium would need a per-user plan\n"
  "  Security and governance — who owns the responses, which account the flow runs as, any DLP policy in force\n"
  "  Performance — expected row counts, delegation support of the chosen data source, connector throttling\n"
  "  Maintainability — who supports the solution afterwards and how a failure would be noticed\n"
  "  A total score and a documented Build / Defer / Reject decision with a one-line justification.\n"
  "Part C — the app created in the course environment (environment picker checked FIRST), named clearly,\n"
  "  with the icon, colour scheme and text customised; screenshot of App settings and of the running app.\n"
  "Evidence: the register, the scored matrix with the decision, and screenshots of the customised app."),

 ("Task 2", "A3, A4",
  "Integrate the app with a data source and support the API-level integration. "
  "Part A — Connect your Employee Engagement Survey app to a cloud data source (OneDrive for Business, "
  "SharePoint or Dataverse). Ensure the app can both read from and write to it. "
  "Part B — Modify the survey questions so they are relevant to your workplace, and bind the input "
  "controls to the data source so a submitted response is stored. Show the Power Fx formulas you used. "
  "Part C — Explain how the app reaches the data source at API level: name the connector, state whether "
  "it is Standard or Premium, and describe what the connector does on the app's behalf. "
  "(Lab 4 — Log to Excel; Lab 7 — Canvas App from Excel Data; Lab 8 — Blank Canvas App with Power Fx; "
  "Lab 9 — Custom Connector and API Integration.)",
  BOX_CAP,
  "Part A — the workbook uploaded to OneDrive for Business and the data formatted as a TABLE\n"
  "  (Insert > Table, named e.g. SurveyResponses). Power Apps cannot bind to a loose range — this is the\n"
  "  most common failure. Data > Add data > Excel Online (Business) > select workbook and table.\n"
  "Part B — controls added and bound, with formulas such as:\n"
  "  Gallery Items:      SortByColumns(Search(SurveyResponses, txtSearch.Text, \"Department\"), \"Date\", Descending)\n"
  "  Dropdown Items:     [\"Engineering\",\"Operations\",\"Sales\",\"Support\",\"Finance\"]\n"
  "  Submit OnSelect:    SubmitForm(EditForm1)   or   Patch(SurveyResponses, Defaults(SurveyResponses),\n"
  "                        {Date: Text(Today()), Department: ddDept.Selected.Value,\n"
  "                         Satisfaction: Value(slSat.Value), Comments: txtComments.Text})\n"
  "  Validation:         If(IsBlank(ddDept.Selected.Value), DisplayMode.Disabled, DisplayMode.Edit)\n"
  "  Preview with F5, submit a response, and confirm the new row appears in the workbook.\n"
  "Part C — API-level explanation: the app never touches the file directly; it calls the Excel Online\n"
  "  (Business) CONNECTOR, which is a packaged REST API. The connector authenticates as the signed-in\n"
  "  user, translates the app's request into Graph/Excel API calls, and returns typed rows the app binds\n"
  "  to. Excel Online (Business), SharePoint and Outlook are Standard connectors; a custom connector\n"
  "  (as built in Lab 9) would be used where no prebuilt connector exists, defining host, base URL,\n"
  "  security and the action's request/response schema.\n"
  "Evidence: screenshots of the data source connected, the formulas, and a new row written to the store."),

 ("Task 3", "A5, A6",
  "Automate the notification, then test and verify the integration. "
  "Part A — Add a Power Automate flow to the app, configured to send an e-mail notification to your "
  "inbox whenever a user submits the survey. Name the flow clearly and show how the app calls it, "
  "including the parameters passed. "
  "Part B — Save and publish the app. "
  "Part C — Perform tests and checks on the connection between the app, the flow and the data source: "
  "submit a test response and verify that the data is stored, the flow run succeeded, and the e-mail "
  "arrived. Show the evidence from the flow's run history. "
  "Part D — Verify the app functions across platforms — in the browser and on the Power Apps mobile "
  "player or embedded in Microsoft Teams — and record the result of each. "
  "(Lab 3 — Trigger and Actions; Lab 11 — Call a Flow from a Canvas App; Lab 12 — Return Data from a "
  "Flow to the App; Lab 10 — Cross-Platform Verification.)",
  BOX_CAP,
  "Part A — flow created with the Power Apps (V2) trigger and TYPED inputs (text/number), e.g.\n"
  "  Applicant (text), Department (text), Satisfaction (number), Comments (text);\n"
  "  action 'Send an email (V2)' (Office 365 Outlook) composing the body from the trigger inputs;\n"
  "  optionally 'Add a row into a table' to log the response.\n"
  "  In the app: Power Automate pane > Add flow > select the flow, then call it from the submit button:\n"
  "    'SurveyNotification'.Run(txtName.Text, ddDept.Selected.Value, Value(slSat.Value), txtComments.Text)\n"
  "  Parameters must be passed in the trigger's declared ORDER and TYPE.\n"
  "Part B — File > Save, then Publish; note that only a published version is served to users, and every\n"
  "  publish creates a restorable version under Details > Version history.\n"
  "Part C — testing evidence:\n"
  "  submit a test response from the app;\n"
  "  open the flow's 28-day RUN HISTORY and show one Succeeded run;\n"
  "  expand each action to show its raw inputs and outputs — proving the data crossed from app to flow;\n"
  "  show the notification e-mail in the inbox;\n"
  "  show the new row in the data source.\n"
  "  Flow checker run with no errors or warnings before saving.\n"
  "Part D — cross-platform verification matrix with a row per app and a column per platform\n"
  "  (Browser / Mobile / Teams), each marked pass or fail with a note — e.g. a Tablet-layout app is\n"
  "  cramped on a phone; a shared app fails for a colleague who lacks access to the data source.\n"
  "Evidence: run-history screenshots, the received e-mail, the stored row, and the completed matrix."),

 ("Task 4", "A7, A8",
  "Highlight the issues and implement the modifications that mitigate them. "
  "Part A — Share the published app with your assessor for testing and review, and capture snapshots of "
  "the key screens including the survey form and the flow configuration. "
  "Part B — During testing, identify at least THREE issues arising from the integration — one technical, "
  "one compatibility and one performance. For each, state the symptom, the evidence you used to "
  "diagnose it, and the root cause. "
  "Part C — Implement a modification that mitigates each issue, and re-test to prove the fix worked. "
  "Record the before and after evidence. "
  "(Lab 14 — Troubleshoot and Optimise the Integration; Lab 13 — End-to-End Approval App.)",
  BOX_CAP,
  "Part A — Share > add the assessor as a User (not Co-owner) so they can run but not edit;\n"
  "  confirm the assessor also has access to the underlying DATA SOURCE, since sharing an app does not\n"
  "  share its data — a finding worth recording in its own right.\n"
  "  Screenshots of the survey form, the confirmation screen and the flow in the designer.\n"
  "Part B/C — an Integration Issue Log with at least three rows, one of each class:\n"
  "  TECHNICAL — symptom: the flow run fails with 401/403 or the app cannot load its data source.\n"
  "    Evidence: the failed run in run history, showing the error on the connector action.\n"
  "    Root cause: an expired, revoked or unauthenticated connection (or a DLP policy, or a missing\n"
  "    Premium licence).\n"
  "    Fix: open the flow, re-authenticate the connection, save and re-run until Succeeded. For a\n"
  "    production solution, run it under a dedicated service account.\n"
  "  COMPATIBILITY — symptom: the call fails schema validation, or a value lands in the wrong field.\n"
  "    Evidence: the schema/type error returned to the app when Days or Satisfaction is passed as text.\n"
  "    Root cause: a type mismatch — text supplied where the trigger declares a number; or a date format\n"
  "    the target cannot parse; or a renamed column.\n"
  "    Fix: coerce the type at the CALL SITE — Value(txtDays.Text) rather than txtDays.Text — and treat\n"
  "    connector inputs as a typed contract. Re-run to prove the call now validates.\n"
  "  PERFORMANCE — symptom: the blue delegation warning on the formula bar; only the first rows appear;\n"
  "    slow load or a throttling error.\n"
  "    Evidence: the delegation warning triangle, and a gallery that returns an incomplete result set.\n"
  "    Root cause: a non-delegable function (e.g. Len(), or a complex nested condition) so only the\n"
  "    first 500 rows are retrieved and filtered on the device.\n"
  "    Fix: rewrite to a delegable formula, e.g.\n"
  "      Filter(SurveyResponses, StartsWith(Department, txtSearch.Text))\n"
  "    until the warning clears; reduce rows at source. Raising the data row limit in Settings > General\n"
  "    is a mitigation, NOT a cure — the candidate should say so.\n"
  "  Each row records: symptom, evidence, classification, root cause, fix implemented, and how the fix\n"
  "  was verified.\n"
  "Evidence: before/after screenshots for each issue, and a final end-to-end run proving the repaired\n"
  "  solution still works."),
]

# ---------------------------------------------------------------- doc helpers
def base_doc():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = "Arial"; n.font.size = Pt(11)
    return doc

def para(doc, text, size=11, bold=False, italic=False, color=None, after=6, before=0, align=None):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(after); p.paragraph_format.space_before = Pt(before)
    if align is not None: p.alignment = align
    return p

def heading(doc, text, size=13):
    para(doc, text, size=size, bold=True, color=BRAND, after=6, before=8)

def answer_box(doc, lines=None, code=None, height_pt=90):
    """1x1 bordered box. `lines` → bullet-style model answer; `code` → monospace
    code/YAML/command block (indentation preserved); neither → empty answer space."""
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.paragraphs[0].text = ""
    if code:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in code.split("\n"):
            b = cell.add_paragraph(style=None)
            b.paragraph_format.space_after = Pt(0); b.paragraph_format.space_before = Pt(0)
            rr = b.add_run(ln if ln else " ")
            rr.font.name = "Consolas"; rr.font.size = Pt(9)
            rr._element.rPr.rFonts.set(qn('w:cs'), "Consolas")
            wt = rr._element.find(qn('w:t'))
            if wt is not None: wt.set(qn('xml:space'), 'preserve')
    elif lines:
        run = cell.paragraphs[0].add_run("Suggestive answers (not exhaustive):")
        run.bold = True; run.font.size = Pt(10.5)
        for ln in lines:
            b = cell.add_paragraph(style=None); b.paragraph_format.left_indent = Inches(0.15)
            rr = b.add_run("•  " + ln); rr.font.size = Pt(10.5)
    else:
        # empty answer space
        tr = t.rows[0]._tr
        trPr = tr.get_or_add_trPr(); trh = OxmlElement('w:trHeight')
        trh.set(qn('w:val'), str(int(height_pt*20))); trh.set(qn('w:hRule'), 'atLeast'); trPr.append(trh)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

FILL_GAP = 6    # extra space below each fill-in line (paired with double line spacing for writing room)

def candidate_block(doc):
    heading(doc, "Trainee Information")
    for label in ["Trainee Name (as per NRIC): ______________________________________",
                  "Last 3 digits and alphabet of NRIC/FIN: ____________________",
                  "Date: ____________________"]:
        p = para(doc, label, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

# Assessment briefing (from the course slides — "Briefing for Assessment").
BRIEFING = [
    "Place phones and other materials under the table or on the floor.",
    "No photos or recording of assessment scripts.",
    "No discussion during the assessment.",
    "Use a black/blue pen for hard-copy assessments.",
    "No liquid paper / correction tape.",
    "Scripts are collected when time is up.",
]

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"

def add_hyperlink(p, url, text):
    """Add a real clickable Word hyperlink (blue, underlined) to paragraph p."""
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)  # 11pt
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rPr.append(color)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link

def instructions(doc, minutes_text):
    heading(doc, "Instructions to Candidate")
    # None marks the upload instruction, which carries a clickable LMS hyperlink.
    items = [
        "This is an individual exercise.",
        "This is an open-book assessment.",
        f"A total of {minutes_text} is given to complete this assessment.",
        None,
    ] + BRIEFING
    for i, s in enumerate(items, 1):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
        if s is None:
            p.add_run(f"{i}.  Complete your answers on the document provided and "
                      "upload the completed answers to the LMS at ").font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)

def grading(doc, what):
    heading(doc, "Grading")
    para(doc, what, size=11, after=12)
    for ln in ["Grade: _______  (C / NYC)",
               "Assessor Name: __________________________   Assessor NRIC: ________________",
               "Date: ________________________                    Signature: ____________________"]:
        p = para(doc, ln, size=11, after=FILL_GAP)
        p.paragraph_format.line_spacing = 2.0

def finish(doc, path):
    prodoc.add_page_numbers(doc); prodoc.enable_update_fields(doc)
    doc.save(path); print("  saved:", os.path.basename(path))

# ---------------------------------------------------------------- builders
def build_wa(answers):
    doc = base_doc()
    kind = "Written Assessment (SAQ) — Answer Key" if answers else "Written Assessment (SAQ)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Written Assessment (SAQ)" if answers else "Written Assessment (SAQ)",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; questions begin on the next page.
        candidate_block(doc); instructions(doc, "1 hour")
        grading(doc, "Candidate has answered all written questions and demonstrated the underpinning "
                     "knowledge required for the course learning outcomes.")
        page_break(doc)
    para(doc, "Short-Answer Questions (Knowledge)", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Answer all questions in your own words. Each question tests underpinning knowledge covered in the "
              "course slides.", size=10.5, italic=True, color=GREY, after=8)
    # Pagination is EXPLICIT — two questions to a page on the paper, one model answer to a
    # page in the key. Do not swap this for Word's keepNext/cantSplit: Word pushes an
    # oversized box to the next page, but Google Docs draws the border anyway and prints the
    # question text and the page footer straight THROUGH it. See SKILL.md → Pagination.
    per_page = 1 if answers else 2
    for i, (crit, ctx, q, pts) in enumerate(WRITTEN, 1):
        para(doc, f"Question {i}:", size=11.5, bold=True, after=2, before=6)
        para(doc, ctx, size=11, after=3)
        para(doc, f"{q}  ({crit})", size=11, bold=True, after=4)
        answer_box(doc, lines=pts if answers else None)
        if i % per_page == 0 and i < len(WRITTEN):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to WA (SAQ) - {TITLE} - {suffix}.docx" if answers
            else f"WA (SAQ) - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

def build_pp(answers):
    doc = base_doc()
    kind = "Practical Performance (PP) — Answer Key" if answers else "Practical Performance (PP)"
    prodoc.add_cover_page(doc, kind, TITLE, A_VER if answers else Q_VER,
                          org_logo=ORG_LOGO, course_logo=COURSE_LOGO)
    para(doc, TITLE, size=15, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, "Answers to Practical Performance Assessment" if answers else "Practical Performance Assessment",
         size=13, bold=True, color=BRAND, align=WD_ALIGN_PARAGRAPH.CENTER, after=2)
    para(doc, f"Course Code: {COURSE_CODE}", size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    if not answers:
        # Page 2 — candidate information, instructions and grading; the problem begins on the next page.
        candidate_block(doc); instructions(doc, "90 minutes")
        grading(doc, "Candidate has successfully completed all PP tasks and can explain the overall "
                     "functions and features used to achieve them.")
        page_break(doc)
    para(doc, "Practical Problem", size=13, bold=True, color=BRAND, after=4)
    para(doc, "Scenario", size=11.5, bold=True, after=2)
    para(doc, SCENARIO, size=11, after=8)
    # Practical tasks are long and their boxes are tall, so they get a page each — on the
    # paper AND in the key. Same rule as the WA: the page break is ours, not the renderer's.
    for i, (label, crit, prompt, cap, pts) in enumerate(PRACTICAL, 1):
        para(doc, f"{label} ({crit}):", size=11.5, bold=True, after=2, before=6)
        para(doc, prompt, size=11, after=3)
        para(doc, cap, size=10.5, italic=True, color=GREY, after=4)
        answer_box(doc, code=pts if answers else None, height_pt=150)
        if i < len(PRACTICAL):
            page_break(doc)
    suffix = A_VER if answers else Q_VER
    name = (f"Answer to PP Assessment - {TITLE} - {suffix}.docx" if answers
            else f"PP Assessment - {TITLE} - {suffix}.docx")
    finish(doc, os.path.join(OUT, name))

if __name__ == "__main__":
    print("Building WSQ assessment set…")
    build_wa(answers=False); build_wa(answers=True)
    build_pp(answers=False); build_pp(answers=True)
    print(f"Done. WA: {len(WRITTEN)} questions · PP: {len(PRACTICAL)} tasks.")
