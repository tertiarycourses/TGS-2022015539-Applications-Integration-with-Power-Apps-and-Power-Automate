#!/usr/bin/env python3
"""Generate the Applications Integration with Power Apps and Power Automate Lesson Plan.

Cover page + Document Version Control Record + auto TOC + Arial 11pt body +
colour-coded 2-day schedule tables (9:30am-6:30pm, 8 hours/day, 1h lunch, tea
within). Timing follows the accredited course proposal: 13.5 hours delivery
(7.5h classroom + 6h practical) plus 2.5 hours assessment (WA 1h + PP 1.5h)
= 16 hours total. Topics/labs come from course_data + the domain data files so
the LP stays aligned with the deck, guide and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
HEADER_FILL="1F6FEB"; TOPIC_FILL="E8F0FE"; BREAK_FILL="FFF4E5"; LUNCH_FILL="FDE9D9"; ASSESS_FILL="E8F7EE"

def lab_titles(nums):
    return "; ".join(f"Lab {a['num']}: {a['title']}" for a in ACT if a['num'] in nums)

# ------------------------------------------------ schedule (single source of truth for timing)
# (start, end, minutes, kind, activity_text)  kind: admin/topic/lab/break/lunch/assess/recap
SCHEDULE = {
 1: (C.DAY_THEMES[1], [
    ("9:30","10:00",30,"admin","Welcome, trainer and learner introductions, course outline, learning outcomes, ground rules and mandatory digital attendance (AM)"),
    ("10:00","11:00",60,"topic","Topic 1 — Opportunities for Using Power Platform Apps: what application integration is, types of middleware and their features (K1), the Power Platform family, and identifying opportunities to connect devices, databases, software and applications (A1)"),
    ("11:00","11:15",15,"break","Tea break (within training time)"),
    ("11:15","12:15",60,"topic","Topic 1 continued — performing a feasibility scan to identify potential middleware (A2): connector availability, data readiness, licensing, security and governance, performance and maintainability"),
    ("12:15","13:00",45,"lab","Hands-on: "+lab_titles([1])),
    ("13:00","14:00",60,"lunch","Lunch break"),
    ("14:00","14:45",45,"lab","Digital attendance (PM). Hands-on: "+lab_titles([2])),
    ("14:45","15:45",60,"topic","Topic 2 — Power Automate: cloud flow types, the trigger-action execution model, connectors and dynamic content, proper usage of middleware (K2) with trainer demonstration"),
    ("15:45","16:00",15,"break","Tea break (within training time)"),
    ("16:00","18:15",135,"lab","Hands-on: "+lab_titles([3,4,5,6])),
    ("18:15","18:30",15,"recap","Day 1 recap, Q&A and end-of-day digital attendance"),
 ]),
 2: (C.DAY_THEMES[2], [
    ("9:30","9:45",15,"recap","Day 1 recap and mandatory digital attendance (AM)"),
    ("9:45","10:30",45,"topic","Topic 3 — Power Apps: canvas versus model-driven apps, the platforms apps run on (K3), data sources and connectors, Power Fx, and the functions of Application Programming Interfaces (K5) with trainer demonstration"),
    ("10:30","10:45",15,"break","Tea break (within training time)"),
    ("10:45","11:35",50,"lab","Hands-on: "+lab_titles([7,8])),
    ("11:35","12:15",40,"lab","Hands-on: "+lab_titles([9,10])),
    ("12:15","13:15",60,"lunch","Lunch break"),
    ("13:15","13:50",35,"topic","Digital attendance (PM). Topic 4 — Integrate Power Apps and Power Automate: calling a flow from a canvas app, returning data to close the round trip, and the technical, compatibility and performance issues that arise after integration (K4)"),
    ("13:50","14:35",45,"lab","Hands-on: "+lab_titles([11,12])),
    ("14:35","14:50",15,"break","Tea break (within training time)"),
    ("14:50","15:40",50,"lab","Hands-on: "+lab_titles([13,14])),
    ("15:40","15:50",10,"recap","Course revision, TRAQOM course feedback survey and Q&A"),
    ("15:50","16:00",10,"assess","Briefing for Assessment and Assessment digital attendance"),
    ("16:00","17:00",60,"assess","Written Assessment (WA) — Short-Answer Questions (SAQ), 1 hour, individual, open book. Five questions covering K1–K5"),
    ("17:00","18:30",90,"assess","Practical Performance (PP) — scenario-based hands-on tasks, 1.5 hours, individual, open book. Four tasks covering A1–A8"),
 ]),
}

# ------------------------------------------------ build document
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)

prodoc.add_cover_page(doc,"LESSON PLAN",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("7.0","1 November 2025","Previous release — Learner Guide slides only, no separate lesson plan schedule.","Tertiary Infotech Academy"),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,
  "Major revision. Rebuilt on the accredited course proposal (13.5h delivery + 2.5h assessment). "
  "Added a progressive 14-lab programme delivered in a dedicated Power Platform environment, "
  "with every K and A statement mapped to a lab. Enhanced with canvas app and Power Automate content.",
  C.TRAINER),
])
prodoc.add_toc(doc)

def H(text,level=1):
    h=doc.add_heading(text,level=level); return h

H("Course Information",1)
info=[("Course Title",C.TITLE),("WSQ Course Reference",C.COURSE_CODE),
      ("TSC Title / Code",f"{C.TSC_TITLE}  ·  {C.TSC_CODE}  ·  Proficiency {C.TSC_LEVEL}"),
      ("Training Provider",C.ORG+"  ("+C.UEN.replace('UEN: ','UEN ')+")"),
      ("Duration","2 days · 16 hours total (13.5 hours delivery + 2.5 hours assessment), 8 hours per day"),
      ("Delivery Breakdown","Classroom facilitation 7.5 hours · Practical / practicum 6 hours"),
      ("Daily Timing","Day 1: 9:30 am – 6:30 pm. Day 2: 9:30 am – 6:30 pm, the final 2.5 hours being the assessment. One-hour lunch; tea breaks within training time."),
      ("Mode","Instructor-led classroom facilitation with hands-on practical labs in a dedicated Microsoft Power Platform environment"),
      ("Lab Environment",C.LAB_ENVIRONMENT),
      ("Trainer-to-Learner Ratio","1:3 to 1:20"),
      ("Trainer",C.TRAINER)]
t=doc.add_table(rows=0,cols=2); t.style="Table Grid"
for k,v in info:
    c=t.add_row().cells; c[0].text=""; r=c[0].paragraphs[0].add_run(k); r.bold=True; r.font.size=Pt(10)
    prodoc._shade_cell(c[0],TOPIC_FILL)
    c[1].text=""; c[1].paragraphs[0].add_run(v).font.size=Pt(10)

H("Learning Outcomes",1)
doc.add_paragraph("On completion of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(lo).font.size=Pt(11)

H("Assessment",1)
for a in [C.ASSESSMENT["written"],C.ASSESSMENT["practical"],
          "Format: Open Book — course slides, Learner Guide and approved materials only.",
          "Assessor-to-candidate ratio: 1:3 to 1:20 for both WA(SAQ) and PP.",
          "The final assessment is conducted at the end of Day 2, after the TRAQOM survey and the Assessment digital attendance.",
          "Closure of gaps: an oral clarification of 5 minutes (assessor-to-candidate 1:1) may be conducted for the WA and the PP if there are minor performance gaps. Time spent on clarification is outside the assessment duration.",
          C.ASSESSMENT["note"]]:
    p=doc.add_paragraph(style="List Bullet"); p.add_run(a).font.size=Pt(11)

def set_cell(cell,text,bold=False,size=9.5,color=None,fill=None,align=None):
    cell.text=""; p=cell.paragraphs[0]
    if align: p.alignment=align
    r=p.add_run(text); r.bold=bold; r.font.size=Pt(size); r.font.name="Arial"
    if color: r.font.color.rgb=color
    if fill: prodoc._shade_cell(cell,fill)

KIND_FILL={"topic":TOPIC_FILL,"break":BREAK_FILL,"lunch":LUNCH_FILL,"assess":ASSESS_FILL,
           "admin":"F3F5F8","recap":"F3F5F8","lab":None}

H("Course Schedule",1)
for day,(theme,rows) in SCHEDULE.items():
    H(f"Day {day} — {theme}",2)
    tbl=doc.add_table(rows=0,cols=3); tbl.style="Table Grid"; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
    hdr=tbl.add_row().cells
    for i,htext in enumerate(["Time","Duration","Topic / Activity"]):
        set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
    deliver=0; assess=0
    for start,end,mins,kind,text in rows:
        cells=tbl.add_row().cells; fill=KIND_FILL.get(kind)
        set_cell(cells[0],f"{start}–{end}",bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        set_cell(cells[1],f"{mins} min",size=9.5,fill=fill)
        set_cell(cells[2],text,bold=(kind in ("topic","assess")),size=9.5,fill=fill)
        if kind=="assess": assess+=mins
        elif kind!="lunch": deliver+=mins
    # widths
    for row in tbl.rows:
        row.cells[0].width=Inches(1.15); row.cells[1].width=Inches(0.9); row.cells[2].width=Inches(4.75)
    p=doc.add_paragraph()
    r=p.add_run(f"Day {day}: {deliver} minutes ({deliver/60:.1f} hours) of delivery"
                + (f" plus {assess} minutes ({assess/60:.1f} hours) of assessment" if assess else "")
                + f" = {deliver+assess} minutes ({(deliver+assess)/60:.1f} hours), excluding the 1-hour lunch break.")
    r.italic=True; r.font.size=Pt(9.5); r.font.color.rgb=GREY
    assert deliver+assess==480, f"Day {day} totals {deliver+assess} min, expected 480"

H("Lab Reference (aligned to the TSC statements)",1)
tt=doc.add_table(rows=0,cols=3); tt.style="Table Grid"
hdr=tt.add_row().cells
for i,htext in enumerate(["Topic / Learning Unit","Maps to","Labs"]):
    set_cell(hdr[i],htext,bold=True,size=10,color=RGBColor(0xFF,0xFF,0xFF),fill=HEADER_FILL)
for tp in C.TOPICS:
    acts=[a for a in ACT if a["topic"]==tp["num"]]
    cells=tt.add_row().cells
    set_cell(cells[0],f"Topic {tp['code']}: {tp['title']}",bold=True,size=9.5,fill=TOPIC_FILL)
    set_cell(cells[1],tp["weighting"],size=9.5,fill=TOPIC_FILL)
    set_cell(cells[2],", ".join(f"Lab {a['num']}" for a in acts),size=9.5)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
OUT=os.path.join(REPO,"courseware",f"LP-{C.SHORT_TITLE}.docx")
doc.save(OUT)
print("Saved",OUT)
