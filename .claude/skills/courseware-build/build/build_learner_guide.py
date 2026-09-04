#!/usr/bin/env python3
"""Generate the Applications Integration with Power Apps and Power Automate Learner Guide as BOTH a
Markdown mirror (LG-*.md at repo root) and a DOCX (LG-*.docx) from one source, so they never diverge.

House format: cover page, Document Version Control Record, auto TOC, Arial 11pt
body, one section per lab (Objective · Goal · What you'll build · Step-by-step
· Test it), plus setup, revision guidance and a glossary. The DETAILED
step-by-step lives here in the LG, never on the slides. All content is
driven by course_data + the domain data files, keeping the LG 100% aligned with
the slide deck, Lesson Plan and labs.
"""
import os, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE=os.path.dirname(os.path.abspath(__file__)); sys.path.insert(0,HERE)
import course_data as C
from data_domain1 import DOMAIN1; from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3; from data_domain4 import DOMAIN4
ACT=DOMAIN1+DOMAIN2+DOMAIN3+DOMAIN4
import prodoc
def _find_repo(start):
    env=os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env): return env
    d=start
    for _ in range(8):
        d=os.path.dirname(d)
        if os.path.isdir(os.path.join(d,"courseware")) and os.path.isdir(os.path.join(d,"labs")): return d
    return os.path.dirname(os.path.dirname(HERE))
REPO=_find_repo(HERE); ASSETS=os.path.join(os.path.dirname(HERE),"assets")

# ---------------- block DSL (single content stream → MD + DOCX) ----------------
B=[]
def h1(t): B.append(("h1",t))
def h2(t): B.append(("h2",t))
def h3(t): B.append(("h3",t))
def p(t):  B.append(("p",t))
def bullets(xs): B.append(("bullets",xs))
def steps(xs): B.append(("steps",xs))
def code(t): B.append(("code",t))
def note(t): B.append(("note",t))
def rule(): B.append(("rule",))
def image(path, caption=""): B.append(("image", path, caption))


def _lab_shot(num):
    """The workflow screenshot for a lab, if one has been captured."""
    import glob as _g
    hits = sorted(_g.glob(os.path.join(ASSETS, "labs", f"lab-{num:02d}-*.png")))
    return hits[0] if hits else None

# ---------------- content ----------------
h1("Introduction")
p(f"This Learner Guide accompanies the WSQ course {C.TITLE} ({C.COURSE_CODE}), conducted by {C.ORG}. "
  f"It maps to the Skills Framework TSC {C.TSC_TITLE} ({C.TSC_CODE}) at Proficiency {C.TSC_LEVEL}, and "
  "provides the full step-by-step instructions for all 14 hands-on labs, organised by the four learning "
  "units of the course.")
p("The labs are PROGRESSIVE: each one builds on the artefact the previous lab produced. Labs 1-2 decide "
  "WHAT to integrate and whether it is feasible; Labs 3-6 build the automation in Power Automate; Labs "
  "7-10 build the apps in Power Apps; and Labs 11-14 join the two halves into one working integration, "
  "then deliberately break, diagnose and repair it. Work them in order — skipping ahead leaves you "
  "without the flow or the table the next lab expects.")
p("Use this guide alongside the course slides and the lab folders in the labs/ directory. The slides give "
  "you the concepts and the shape of each lab; this guide gives you the detailed procedure.")

h1("Course Learning Outcomes")
bullets(C.LEARNING_OUTCOMES)

h1("Skills Framework Mapping")
p("Every knowledge (K) and ability (A) statement in the TSC is taught and assessed. The Written "
  "Assessment covers the K statements; the Practical Performance covers the A statements.")
h3("Knowledge statements — assessed by the Written Assessment (SAQ)")
bullets([f"{k} — {d}" for k, d in C.TSC_KNOWLEDGE])
h3("Ability statements — assessed by the Practical Performance (PP)")
bullets([f"{a} — {d}" for a, d in C.TSC_ABILITIES])

h1("Before You Start — Environment Setup")
h3("What you need")
bullets([
 "A Microsoft 365 work or school account with a Power Platform licence — your trainer issues the training account for the class.",
 f"Access to the Power Platform environment '{C.LAB_ENVIRONMENT}', which has been provisioned for this course with Microsoft Dataverse enabled.",
 "A modern browser (Microsoft Edge or Google Chrome). Everything in this course runs in the browser — nothing is installed on your machine.",
 "OneDrive for Business, which stores the Excel workbooks the flows and apps read and write.",
 "Microsoft Teams, because approval requests are delivered to the Teams Approvals app.",
 "The lab data workbooks, downloaded from the LMS. Each lab folder contains the workbook that lab needs.",
])
h3("Switch to the course environment FIRST")
p("This is the single most important setup step. Power Platform keeps apps and flows inside an "
  "environment, and work created in the wrong environment cannot simply be dragged into the right one. "
  "Before you build anything, check the environment picker in the top-right corner of the maker portal.")
steps([
 ("Open the Power Automate maker portal.", "https://make.powerautomate.com"),
 ("Click the environment picker in the top-right corner of the page.", ""),
 (f"Select '{C.LAB_ENVIRONMENT}'.", ""),
 ("Confirm the environment name is now shown in the top-right before you continue.", ""),
 ("Repeat the same check at make.powerapps.com — the two portals track the environment separately.",
  "https://make.powerapps.com"),
])
h3("Upload the lab data to OneDrive")
p("Power Apps and Power Automate can only bind to a real Excel TABLE, never to a loose range of cells. "
  "The workbooks supplied with each lab already contain a correctly named table, so upload them "
  "unchanged. If you later build your own workbook, remember to select your data and choose "
  "Insert > Table, then give the table a name in Table Design — forgetting this is the single most "
  "common reason a data source fails to appear.")
steps([
 ("Open OneDrive for Business in your browser and sign in with your training account.",
  "https://www.office.com/launch/onedrive"),
 ("Create a folder named PowerPlatformLabs at the root of your OneDrive.", ""),
 ("Upload the workbook from each lab's data/ folder into it — 'KinetEco Service Calls.xlsx', 'LeaveLog.xlsx' and 'Employee Survey.xlsx'.", ""),
 ("Open one workbook and confirm the table name under Table Design — ServiceCalls, LeaveLog and SurveyResponses respectively.", ""),
])
h3("Naming conventions used in every lab")
bullets([
 "Name every flow and app EXACTLY as the lab specifies, including the trailing (DO NOT DELETE).",
 "The (DO NOT DELETE) suffix marks the object as courseware in a shared training tenant, so that housekeeping does not remove another learner's work.",
 "Formulas shown in a code block are Power Fx (for apps) or Power Automate expressions (for flows) — type them into the formula bar or the expression editor, not into a terminal.",
 "Placeholders such as <your-email> are replaced with your own values.",
 "If a lab tells you to break something on purpose, do it — provoking a fault you caused is the fastest way to learn to diagnose one you did not.",
])
h3("Using the prebuilt packages")
p("Every lab folder ships the finished flow in two formats, so you can inspect or restore a working "
  "version at any point. Building it yourself is always the better learning route; import only if you "
  "fall behind or want to compare your build against the reference.")
bullets([
 "Solution-Lab-NN.zip — a Dataverse solution. Import via Solutions > Import solution. This is the format that imports on the course tenant.",
 "Lab-N-....zip — a legacy package. Import via My flows > Import > Import Package (Legacy). Use this on tenants that do not have 'Create in Dataverse solutions' enabled.",
 "labs/Solution-All-Labs.zip installs every lab flow at once.",
 "Imported flows arrive TURNED OFF and without connections. Open each one, re-authenticate every connector, then turn it on — this is itself the technical-issue pattern you diagnose in Lab 14.",
])

# ---------------- per-topic, per-lab ----------------
TOPICS_BY_NUM={t["num"]:t for t in C.TOPICS}
for t in C.TOPICS:
    h1(f"Topic {t['code']} — {t['title']}")
    p(t["subtitle"])
    h3("Key concepts")
    bullets(t["concepts"])
    for a in [x for x in ACT if x["topic"]==t["num"]]:
        h2(f"Lab {a['num']} — {a['title']}")
        p(f"Maps to: {a['objective']}")
        p(f"Goal: {a['desc']}")
        h3("What you'll build")
        p(a["build"]+f"   (Tools: {a['services']}.)")
        _shot=_lab_shot(a["num"])
        if _shot:
            image(_shot, f"Lab {a['num']} — {a['title']}: the workflow in the course environment.")
        h3("Step-by-step")
        st=[]
        for i,(instr,cmd) in enumerate(a["steps"],1):
            st.append((instr,cmd))
        steps(st)
        h3("Test it")
        p(a["test"])
        note(f"This lab has its own folder at labs/lab-{a['num']:02d}/ containing the lab sheet, the "
             f"data workbook it needs and the prebuilt flow packages. Build every object inside the "
             f"course environment '{C.LAB_ENVIRONMENT}'.")
        rule()

h1("Reference — Integration Issues and How to Fix Them")
p("Integration defects fall into three classes. The Practical Performance assessment asks you to "
  "highlight issues following integration (A7) and implement modifications to mitigate them (A8), so "
  "learn to name the class, cite the evidence and state the fix.")
h3("Technical issues — the connection itself")
bullets([
 "Symptoms: the flow run fails with a 401 or 403; the connection shows a warning triangle; the app cannot load its data source.",
 "Causes: an expired, revoked or never-authenticated connection; the wrong account; a Data Loss Prevention policy separating two connectors you tried to combine; a missing Premium licence.",
 "Fixes: open the flow, re-authenticate each connection, and save. For anything that must outlive one person, run it under a dedicated service account rather than a personal account. Check the environment's DLP policies before designing a flow that spans connector groups.",
])
h3("Compatibility issues — the shape of the data")
bullets([
 "Symptoms: a schema validation error the moment the flow is called; the right value lands in the wrong field; a date arrives as text.",
 "Causes: a type mismatch — text passed where the connector declares a number; a date format the target does not parse; an Excel range that is not a Table; a renamed column.",
 "Fixes: coerce the type at the CALL SITE, for example Value(txtDays.Text) rather than txtDays.Text. Treat a connector's inputs as a typed contract. Keep column names stable, and re-point the action after any rename.",
])
h3("Performance issues — scale and throughput")
bullets([
 "Symptoms: a blue delegation warning in the formula bar; only the first rows appear; the app is slow to load; the connector returns a throttling error.",
 "Causes: a non-delegable function such as Len() or a complex nested condition; too many rows pulled to the device; connector request limits exceeded.",
 "Fixes: rewrite to a delegable function such as StartsWith() or Filter() on an indexed column; reduce rows at source; raise the data row limit in Settings > General as a stop-gap, understanding that it is a mitigation and not a cure; batch or throttle high-volume calls.",
])
h3("Delegation — the one to remember")
p("Delegation means the filtering work is pushed DOWN to the data source rather than done on the "
  "device. When a formula cannot be delegated, Power Apps retrieves only the first 500 rows (raisable "
  "to 2,000) and filters those locally. The app does not fail — it silently shows an incomplete answer, "
  "which is far more dangerous than an error. Watch for the blue warning triangle in the formula bar and "
  "rewrite the formula until it disappears.")
rule()

h1("Revision and Assessment Preparation")
h3("How to revise")
bullets([
 "First pass: complete every lab in order, in the course environment, following this guide.",
 "Second pass: rebuild Labs 3, 4 and 11 from an empty flow without looking, until the trigger-action-test rhythm is automatic.",
 "For each lab, re-read the 'Test it' criterion and satisfy yourself you could demonstrate it to an assessor.",
 "Practise naming the three classes of integration issue and giving one concrete example and fix for each.",
 "Be able to explain, in your own words, what middleware is and why the Power Platform is an example of it.",
 "Practise the knowledge questions at the Tertiary Infotech practice exam site: https://exams.tertiaryinfotech.com",
])
h3("What the Written Assessment (SAQ) covers")
p("Five open-ended short-answer questions, one hour, individual and open book. The questions test the "
  "underpinning knowledge: types of middleware and their features (K1), proper usage of middleware (K2), "
  "the platforms applications run on (K3), the technical, compatibility and performance issues that arise "
  "in integration (K4), and the functions of Application Programming Interfaces (K5).")
h3("What the Practical Performance (PP) covers")
p("Four scenario-based hands-on tasks, ninety minutes, individual and open book. You build and "
  "demonstrate a working integration: identifying the opportunity and scanning feasibility (A1, A2), "
  "using the platform to integrate data and functions (A3), supporting API-level integration (A4), "
  "testing the connections (A5), verifying the modules work across platforms (A6), then highlighting the "
  "issues you meet (A7) and implementing the modifications that fix them (A8).")
h3("On assessment day")
bullets([
 "Complete the TRAQOM course feedback survey on the LMS.",
 "Take the Assessment digital attendance by scanning the SSG QR code.",
 "Sit the Written Assessment, then the Practical Performance.",
 "Submit your answers on the LMS at https://lms-tms.tertiaryinfotech.com.",
 "Sign the Assessment Summary Record.",
 "A minimum of 75% attendance is required to be eligible for assessment and funding.",
])

h1("Glossary")
gl=[
 ("Application integration","Connecting two or more applications so that data and functions pass between them without manual re-keying."),
 ("Middleware","The software layer that sits between applications and brokers the exchange. Power Platform is low-code integration middleware."),
 ("iPaaS","Integration Platform as a Service — cloud middleware with prebuilt connectors, such as Power Automate."),
 ("Connector","A packaged API. Standard connectors are included in most licences; Premium connectors require a higher plan; a custom connector wraps any REST API."),
 ("Custom connector","A connector you define yourself over an existing REST API by describing its host, security, actions and response schema."),
 ("Trigger","The single event that starts a cloud flow. Every flow has exactly one."),
 ("Action","A step the flow performs after the trigger. Actions run in the order set by their runAfter dependencies."),
 ("Dynamic content","The output of one action, inserted as the input of a later one — how data crosses application boundaries inside a flow."),
 ("Cloud flow","A flow that runs in the Power Automate service. Automated, instant and scheduled are the three types."),
 ("Desktop flow (RPA)","A flow that drives a user interface, used for systems that expose no API."),
 ("Canvas app","A Power App whose interface you design yourself, control by control. It can bind to any connector."),
 ("Model-driven app","A Power App whose interface is generated from a Dataverse data model."),
 ("Power Fx","The Excel-like formula language that binds Power Apps controls to data and behaviour."),
 ("Dataverse","The governed data platform underlying Power Platform, providing tables, relationships and a security model."),
 ("Delegation","Pushing filtering and sorting down to the data source instead of doing it on the device. Non-delegable formulas silently process only the first rows."),
 ("Environment","A container for apps, flows and data. Work created in one environment cannot simply be moved to another."),
 ("Solution","A package of Power Platform components used to move apps and flows between environments."),
 ("Connection reference","A named pointer to a connection, used inside a solution so the same flow can bind to different credentials in each environment."),
 ("DLP policy","A Data Loss Prevention policy that prevents specified connectors being combined in one flow or app."),
 ("Approval","A built-in Power Automate action that requests a decision from a person and waits for the response."),
 ("Run history","The 28-day record of every flow run, with the inputs and outputs of each action — your primary evidence when testing and diagnosing."),
 ("Flow checker","The built-in validator that reports errors and warnings in a flow before you save it."),
 ("Feasibility scan","A structured assessment of whether a candidate integration should be built, weighing connector availability, data readiness, licensing, security, performance and maintainability."),
]
B.append(("dl",gl))

# ---------------- render Markdown ----------------
def _anchor(txt):
    return "".join(ch.lower() if ch.isalnum() else ("-" if ch in " -" else "") for ch in txt)

def render_md():
    out=[f"# {C.TITLE} — Learner Guide",""]
    out.append(f"**WSQ Course Code:** {C.COURSE_CODE}  |  **Conducted by:** {C.ORG} ({C.UEN.replace('UEN: ','UEN ')})  |  **Version {C.VERSION} · {C.VERSION_DATE}**")
    out.append("")
    # TOC (h1 + h2)
    out.append("## Contents"); out.append("")
    for kind,*rest in B:
        if kind=="h1": out.append(f"- [{rest[0]}](#{_anchor(rest[0])})")
        elif kind=="h2": out.append(f"  - [{rest[0]}](#{_anchor(rest[0])})")
    out.append("")
    for kind,*rest in B:
        if kind=="h1": out+=["",f"## {rest[0]}",""]
        elif kind=="h2": out+=["",f"### {rest[0]}",""]
        elif kind=="h3": out+=[f"**{rest[0]}**",""]
        elif kind=="p": out+=[rest[0],""]
        elif kind=="bullets": out+=[f"- {x}" for x in rest[0]]+[""]
        elif kind=="steps":
            for i,(instr,cmd) in enumerate(rest[0],1):
                out.append(f"{i}. {instr}")
                if cmd: out+=["",f"   ```bash",f"   {cmd}","   ```",""]
            out.append("")
        elif kind=="code": out+=["```bash",rest[0],"```",""]
        elif kind=="note": out+=[f"> **Note:** {rest[0]}",""]
        elif kind=="image":
            rel=os.path.relpath(rest[0], REPO)
            out+=[f"![{rest[1]}]({rel})",""]
            if rest[1]: out+=[f"*{rest[1]}*",""]
        elif kind=="rule": out+=["---",""]
        elif kind=="dl":
            for term,defn in rest[0]: out.append(f"- **{term}** — {defn}")
            out.append("")
    return "\n".join(out)

MD_OUT=os.path.join(REPO,f"LG-{C.SHORT_TITLE}.md")
with open(MD_OUT,"w") as f: f.write(render_md())
print("Saved",MD_OUT)

# ---------------- render DOCX ----------------
BRAND=RGBColor(0x1F,0x6F,0xEB); DARK=RGBColor(0x11,0x18,0x27); GREY=RGBColor(0x55,0x5B,0x66)
INKCODE=RGBColor(0x0B,0x30,0x60)
doc=Document()
normal=doc.styles["Normal"]; normal.font.name="Arial"; normal.font.size=Pt(11)
prodoc.style_headings(doc)
prodoc.add_cover_page(doc,"LEARNER GUIDE",C.TITLE,C.VERSION.lstrip("v"),
                      org_logo=os.path.join(ASSETS,"tertiary-infotech-logo.png"),
                      course_logo=None, course_code=C.COURSE_CODE)
prodoc.add_version_control(doc,[
 ("7.0","1 November 2025","Previous release — Learner Guide slides.","Tertiary Infotech Academy"),
 (C.VERSION.lstrip("v"),C.VERSION_DATE,
  "Major revision. Rebuilt from the accredited course proposal with a progressive 14-lab programme "
  "delivered in a dedicated Power Platform environment. Enhanced with canvas app and Power Automate "
  "content; added the Skills Framework mapping, the integration-issue reference (technical, "
  "compatibility, performance) and a full glossary. Every K and A statement is mapped to a lab.",
  C.TRAINER),
])
prodoc.add_toc(doc)

def code_para(text):
    for line in text.split("\n"):
        para=doc.add_paragraph(); prodoc._shade_para(para) if hasattr(prodoc,"_shade_para") else None
        r=para.add_run(line); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=INKCODE


# ---------------- numbered-step restart ----------------
# Word's "List Number" style shares ONE numbering instance across the whole document,
# so every lab's steps continue from the previous lab (Lab 1 would start at 10).
# Each steps block therefore gets its OWN w:num that restarts at 1.
def _new_numbering_id(doc):
    """Create a fresh numbering instance WITH ITS OWN abstractNum and return its numId.

    Instances that share one abstractNum continue each other's count in Word and
    LibreOffice, so each steps block needs a private abstract definition to restart at 1.
    """
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _E
    numbering = doc.part.numbering_part.element

    used_abs = [int(a.get(_qn('w:abstractNumId')))
                for a in numbering.findall(_qn('w:abstractNum'))]
    abs_id = max(used_abs + [8999]) + 1

    abstract = _E('w:abstractNum')
    abstract.set(_qn('w:abstractNumId'), str(abs_id))
    nsid = _E('w:nsid'); nsid.set(_qn('w:val'), f'{abs_id:08X}'); abstract.append(nsid)
    mkind = _E('w:multiLevelType'); mkind.set(_qn('w:val'), 'singleLevel'); abstract.append(mkind)
    lvl = _E('w:lvl'); lvl.set(_qn('w:ilvl'), '0')
    start = _E('w:start'); start.set(_qn('w:val'), '1'); lvl.append(start)
    fmt = _E('w:numFmt'); fmt.set(_qn('w:val'), 'decimal'); lvl.append(fmt)
    rst = _E('w:lvlRestart'); rst.set(_qn('w:val'), '0'); lvl.append(rst)
    txt = _E('w:lvlText'); txt.set(_qn('w:val'), '%1.'); lvl.append(txt)
    jc = _E('w:lvlJc'); jc.set(_qn('w:val'), 'left'); lvl.append(jc)
    ppr = _E('w:pPr')
    ind = _E('w:ind'); ind.set(_qn('w:left'), '425'); ind.set(_qn('w:hanging'), '360')
    ppr.append(ind); lvl.append(ppr)
    abstract.append(lvl)
    # abstractNum elements must precede num elements in numbering.xml
    first_num = numbering.find(_qn('w:num'))
    if first_num is not None:
        first_num.addprevious(abstract)
    else:
        numbering.append(abstract)

    used = [int(n.get(_qn('w:numId'))) for n in numbering.findall(_qn('w:num'))]
    new_id = max(used + [0]) + 1
    num = _E('w:num'); num.set(_qn('w:numId'), str(new_id))
    ref = _E('w:abstractNumId'); ref.set(_qn('w:val'), str(abs_id)); num.append(ref)
    numbering.append(num)
    return new_id


def _numbered(doc, text, num_id):
    """A numbered paragraph bound to the given numbering instance."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _E
    para = doc.add_paragraph(style="List Number")
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(_qn('w:numPr')):
        pPr.remove(old)
    numPr = _E('w:numPr')
    ilvl = _E('w:ilvl'); ilvl.set(_qn('w:val'), '0'); numPr.append(ilvl)
    nid = _E('w:numId'); nid.set(_qn('w:val'), str(num_id)); numPr.append(nid)
    pPr.append(numPr)
    para.add_run(text)
    return para

for kind,*rest in B:
    if kind=="h1": doc.add_heading(rest[0],level=1)
    elif kind=="h2": doc.add_heading(rest[0],level=2)
    elif kind=="h3":
        para=doc.add_paragraph(); r=para.add_run(rest[0]); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=BRAND
    elif kind=="p": doc.add_paragraph(rest[0])
    elif kind=="bullets":
        for x in rest[0]: doc.add_paragraph(x,style="List Bullet")
    elif kind=="steps":
        _nid=_new_numbering_id(doc)          # restart at 1 for THIS block
        for i,(instr,cmd) in enumerate(rest[0],1):
            _numbered(doc,instr,_nid)
            if cmd: code_para(cmd)
    elif kind=="code": code_para(rest[0])
    elif kind=="note":
        para=doc.add_paragraph(); r=para.add_run("Note: "); r.bold=True; r.font.color.rgb=BRAND
        para.add_run(rest[0]).font.size=Pt(10)
    elif kind=="image":
        try:
            from docx.shared import Inches as _In
            para=doc.add_paragraph(); para.alignment=WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(rest[0], width=_In(6.0))
            if rest[1]:
                cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
                r=cap.add_run(rest[1]); r.italic=True; r.font.size=Pt(9); r.font.color.rgb=GREY
        except Exception as e:
            print(f"  [warn] could not embed {rest[0]}: {e}")
    elif kind=="rule": doc.add_paragraph("")
    elif kind=="dl":
        for term,defn in rest[0]:
            para=doc.add_paragraph(style="List Bullet")
            r=para.add_run(term+" — "); r.bold=True; para.add_run(defn)

prodoc.add_page_numbers(doc)
prodoc.enable_update_fields(doc)
DOCX_OUT=os.path.join(REPO,"courseware",f"LG-{C.SHORT_TITLE}.docx")
doc.save(DOCX_OUT)
print("Saved",DOCX_OUT)
